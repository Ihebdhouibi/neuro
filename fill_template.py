#!/usr/bin/env python3
"""
Word Prescription Template Filler

Fills the existing Word (.docx) prescription template with OCR-extracted data
and converts it to PDF.

Template: templates/prescription/Ordonnance_Template vierge.docx
Fields filled from OCR extraction schema (openai_extraction.py):
  - form_date, patient.{last_name, first_name, nir, birth_date}
  - doctor.{full_name, rpps}
  - orthoptic_care.{description, acts_prescribed}
  + center info (finess, address, etc.) from configuration
"""

import os
import copy
import shutil
from pathlib import Path
from typing import Dict, List, Optional
from datetime import datetime

import re

from docx import Document
from docx.shared import Pt, RGBColor
from loguru import logger

# Default template path (relative to this file)
DEFAULT_TEMPLATE = Path(__file__).parent / "templates" / "prescription" / "Ordonnance_Template vierge.docx"

# Official AMY Procedure Table (duplicated from prescription_generator.py to avoid
# importing PyMuPDF-dependent module; keep in sync or move to shared config)
AMY_TABLE: Dict[str, Dict[str, str]] = {
    'AMY 8': {'label': "Mesure de l'acuité visuelle et de la réfraction – Renouvellement", 'price': '20,80 €'},
    'AMY 15': {'label': 'Bilan des troubles oculomoteurs', 'price': '39,00 €'},
    'AMY 7,7': {'label': 'Séance orthoptique (courte)', 'price': '15,40 €'},
    'AMY 7': {'label': "Traitement de l'amblyopie par série de vingt séances", 'price': '18,20 €'},
    'AMY 4': {'label': 'Traitement des hétérophories (20 séances)', 'price': '10,40 €'},
    'AMY 7.7': {'label': 'Traitement du strabisme par série de vingt séances', 'price': '20,02 €'},
    'AMY 30': {'label': 'Bilan orthoptique des déficiences visuelles (basse vision)', 'price': '78,00 €'},
    'AMY 30,5': {'label': 'Bilan des conséquences neuro-ophtalmologiques', 'price': '79,30 €'},
    'AMY 10': {'label': 'Bilan des déséquilibres de la vision binoculaire', 'price': '26,00 €'},
    'AMY 14,5': {'label': 'Bilan des déséquilibres avec trouble neurosensoriel/accommodatif', 'price': '37,70 €'},
    'AMY 15,5': {'label': "Bilan d'une amblyopie", 'price': '40,30 €'},
    'AMY 19,2': {'label': 'Rééducation déficience visuelle - Plus de 16 ans (45 mn)', 'price': '49,92 €'},
    'AMY 13,2': {'label': 'Rééducation déficience visuelle - 16 ans et moins (30 mn)', 'price': '34,32 €'},
    'AMY 8,7': {'label': 'Mesure acuité visuelle et réfraction - Primo-prescription', 'price': '22,62 €'},
}


def normalize_amy_code(raw: str) -> Optional[str]:
    """Normalize AMY code from OCR text (e.g., 'AMY (8)' -> 'AMY 8')."""
    if not raw:
        return None
    match = re.search(r'AMY\s*\(?\s*(\d+(?:[,\.]\d+)?)\s*\)?', raw, re.IGNORECASE)
    if not match:
        return None
    num = match.group(1).replace('.', ',')
    return f'AMY {num}'

# ---------- helpers ----------

def _format_date_fr(iso_date: Optional[str]) -> str:
    """Convert YYYY-MM-DD to DD/MM/YYYY for display. Returns '' if None."""
    if not iso_date:
        return ""
    try:
        dt = datetime.strptime(iso_date, "%Y-%m-%d")
        return dt.strftime("%d/%m/%Y")
    except ValueError:
        return iso_date  # return as-is if already formatted


def _append_value_to_run(paragraph, label_text: str, value: str):
    """
    Find the run containing *label_text* and append *value* right after it
    in a new run that inherits the paragraph's body style (dark grey, normal weight).
    """
    for run in paragraph.runs:
        if label_text in run.text:
            # Create a new run after this one with the value
            new_run = copy.deepcopy(run._element)
            new_run.text = f" {value}"
            # Style: not bold, dark text
            rpr = new_run.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
            if rpr is not None:
                bold = rpr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}b')
                if bold is not None:
                    rpr.remove(bold)
            run._element.addnext(new_run)
            return True
    return False


def _set_date_parts(paragraph, day: str, month: str, year: str):
    """
    Fill the date paragraph which has format:  Date : ____ / ____ / ____
    The runs contain 'Date :', spaces, ' /', spaces, ' / ', spaces.
    We insert day/month/year values into the space runs.
    """
    runs = paragraph.runs
    # Find the run containing 'Date :'
    date_idx = None
    for i, r in enumerate(runs):
        if 'Date :' in r.text:
            date_idx = i
            break
    if date_idx is None:
        return

    # After 'Date :' there are spacing runs, then ' /', then spacing, then ' / '
    # Strategy: replace the first whitespace-only run in each slot with the value,
    # then blank out remaining spacing runs in that slot.
    slash_count = 0
    parts = [day, month, year]
    filled = [False, False, False]
    for i in range(date_idx + 1, len(runs)):
        text = runs[i].text
        if '/' in text:
            slash_count += 1
            continue
        if slash_count < len(parts) and text.strip() == '':
            if not filled[slash_count]:
                runs[i].text = f" {parts[slash_count]} "
                filled[slash_count] = True
            else:
                runs[i].text = ""


def _insert_acts(doc, acts_with_labels: List[Dict[str, str]]):
    """
    Insert prescribed acts into the empty paragraphs after 'ACTES PRESCRITS :' (P[10]).
    Each act gets: "- {code} : {label}  ({price})"
    """
    # Find the ACTES PRESCRITS paragraph index
    actes_idx = None
    for i, p in enumerate(doc.paragraphs):
        if 'ACTES PRESCRITS' in p.text:
            actes_idx = i
            break

    if actes_idx is None:
        return

    # Use paragraphs P[11] through P[20] (up to 10 empty slots)
    for j, act in enumerate(acts_with_labels):
        target_idx = actes_idx + 1 + j
        if target_idx >= len(doc.paragraphs):
            break
        p = doc.paragraphs[target_idx]
        # Clear any existing content
        for run in p.runs:
            run.text = ""
        # Add act text
        run = p.add_run(f"- {act['code']} : {act['label']}    ({act['price']})")
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


# ---------- main fill function ----------

def fill_prescription_template(
    extraction_data: Dict,
    center_info: Optional[Dict[str, str]] = None,
    finess: str = "",
    city: str = "",
    amy_table: Optional[Dict[str, Dict[str, str]]] = None,
    template_path: Optional[str] = None,
) -> str:
    """
    Fill the Word prescription template with extracted OCR data.

    Args:
        extraction_data: OCR extraction result matching sample_schema.json:
            {form_date, patient: {last_name, first_name, nir, birth_date},
             doctor: {full_name, rpps},
             orthoptic_care: {description, acts_prescribed}}
        center_info: Optional dict with keys: name, address, tel, email
        finess: FINESS number of the center
        city: City for "Fait à" field
        amy_table: AMY code → {label, price} lookup. Falls back to prescription_generator.AMY_TABLE
        template_path: Override path to .docx template

    Returns:
        Path to the filled .docx file (in a temp location, ready for PDF conversion)
    """
    tpl = Path(template_path) if template_path else DEFAULT_TEMPLATE
    if not tpl.exists():
        raise FileNotFoundError(f"Template not found: {tpl}")

    doc = Document(str(tpl))

    # --- Unpack extraction data ---
    patient = extraction_data.get("patient", {})
    doctor = extraction_data.get("doctor", {})
    care = extraction_data.get("orthoptic_care", {})
    form_date = _format_date_fr(extraction_data.get("form_date"))

    # --- P[0]: Center name ---
    if center_info and center_info.get("name"):
        for run in doc.paragraphs[0].runs:
            if run.text == "Nom du centre":
                run.text = center_info["name"]
                break

    # --- P[1]: Address / Tel / Email ---
    if center_info:
        p1 = doc.paragraphs[1]
        for run in p1.runs:
            if "Adresse" in run.text:
                addr = center_info.get("address", "")
                run.text = f"Adresse\u00a0: {addr}"
            elif "Tél" in run.text:
                tel = center_info.get("tel", "")
                run.text = f"\nTél : {tel}"
            elif "Email" in run.text:
                email = center_info.get("email", "")
                run.text = f"\nEmail : {email}"

    # --- P[2]: FINESS + Fait à ---
    if finess:
        p2 = doc.paragraphs[2]
        for run in p2.runs:
            if "FINESS :" in run.text:
                run.text = f"FINESS : {finess}"
                break
    if city:
        p2 = doc.paragraphs[2]
        for run in p2.runs:
            if "Fait à" in run.text:
                run.text = f"Fait à {city}"
                break

    # --- P[3]: Date DD / MM / YYYY ---
    if form_date:
        parts = form_date.split("/")
        if len(parts) == 3:
            _set_date_parts(doc.paragraphs[3], parts[0], parts[1], parts[2])

    # --- P[5]: Nom / Prénom ---
    p5 = doc.paragraphs[5]
    last_name = patient.get("last_name") or ""
    first_name = patient.get("first_name") or ""
    for run in p5.runs:
        if run.text == "Nom :":
            run.text = f"Nom : {last_name}"
        elif "Prénom :" in run.text:
            run.text = f"Prénom : {first_name}"

    # --- P[6]: NIR ---
    nir = patient.get("nir") or ""
    p6 = doc.paragraphs[6]
    for run in p6.runs:
        if "Sécurité Sociale" in run.text:
            run.text = f"N° Sécurité Sociale (NIR) : {nir}"
            break

    # --- P[7]: Date de naissance ---
    birth = _format_date_fr(patient.get("birth_date"))
    p7 = doc.paragraphs[7]
    for run in p7.runs:
        if "Date de naissance" in run.text:
            run.text = f"Date de naissance : {birth}"
            break

    # --- P[8]: Docteur / RPPS ---
    doc_name = doctor.get("full_name") or ""
    rpps = doctor.get("rpps") or ""
    p8 = doc.paragraphs[8]
    for run in p8.runs:
        if "Docteur :" in run.text:
            run.text = f"Je soussigné(e), Docteur : {doc_name}"
        elif run.text.strip() == "RPPS":
            # The RPPS label and ':' are in separate runs; put value in the ':' run
            pass
    # Handle RPPS value — find the ' :' run after 'RPPS'
    for i, run in enumerate(p8.runs):
        if run.text.strip() == "RPPS" and i + 1 < len(p8.runs):
            p8.runs[i + 1].text = f" : {rpps}"
            break

    # --- P[9]: Orthoptic care description ---
    # Note: P[9] has runs ['L', 'es soins orthoptiques suivants :'] — clear the
    # leading fragment and write the full text into the main run.
    description = care.get("description") or ""
    if description:
        p9 = doc.paragraphs[9]
        found = False
        for run in p9.runs:
            if "soins orthoptiques" in run.text:
                run.text = f"Les soins orthoptiques suivants : {description}"
                found = True
            elif found is False and run.text.strip() and "soins" not in run.text:
                # Leading fragment (e.g. 'L') before the main run — clear it
                pass
        # Clear any fragment runs before the main label
        for run in p9.runs:
            if run.text.strip() == 'L':
                run.text = ''
                break

    # --- P[11+]: Acts prescribed (AMY codes) ---
    acts_raw = care.get("acts_prescribed", [])
    if acts_raw:
        lookup = amy_table if amy_table is not None else AMY_TABLE

        acts_with_labels = []
        for code in acts_raw:
            normalized = normalize_amy_code(code) or code
            info = lookup.get(normalized, {})
            acts_with_labels.append({
                "code": normalized,
                "label": info.get("label", ""),
                "price": info.get("price", ""),
            })
        _insert_acts(doc, acts_with_labels)

    # --- Save filled document ---
    output_dir = Path(tpl).parent / "filled"
    output_dir.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filled_path = output_dir / f"Ordonnance_filled_{timestamp}.docx"
    doc.save(str(filled_path))
    logger.info(f"Filled template saved: {filled_path}")
    return str(filled_path)


def fill_and_convert_to_pdf(
    extraction_data: Dict,
    output_pdf_path: str,
    center_info: Optional[Dict[str, str]] = None,
    finess: str = "",
    city: str = "",
    amy_table: Optional[Dict[str, Dict[str, str]]] = None,
    template_path: Optional[str] = None,
) -> str:
    """
    Fill the Word template and convert to PDF.

    Args:
        extraction_data: OCR extraction result (see fill_prescription_template)
        output_pdf_path: Desired final PDF path
        center_info: Optional center information dict
        finess: FINESS number
        city: City for "Fait à"
        amy_table: AMY code lookup table
        template_path: Override template .docx path

    Returns:
        Path to the generated PDF
    """
    # Step 1: Fill template
    filled_docx = fill_prescription_template(
        extraction_data=extraction_data,
        center_info=center_info,
        finess=finess,
        city=city,
        amy_table=amy_table,
        template_path=template_path,
    )

    # Step 2: Convert to PDF
    os.makedirs(os.path.dirname(output_pdf_path), exist_ok=True)

    try:
        from docx2pdf import convert
        convert(filled_docx, output_pdf_path)
        logger.info(f"PDF generated: {output_pdf_path}")
    except Exception as e:
        logger.error(f"docx2pdf conversion failed: {e}")
        raise

    # Step 3: Clean up intermediate .docx
    try:
        os.unlink(filled_docx)
    except OSError:
        pass

    return output_pdf_path
