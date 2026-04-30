#!/usr/bin/env python3
"""
Word Prescription Template Filler — content-based lookup (resilient to template edits)

Sections are located by searching for marker text rather than fixed paragraph indices.
This survives template edits in Word that collapse/insert paragraphs.
"""

import os
from pathlib import Path
from typing import Dict, List, Optional
from datetime import datetime
import re
import copy

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from loguru import logger

DEFAULT_TEMPLATE      = Path(__file__).parent / "templates" / "prescription" / "Ordonnance_Template vierge.docx"
DEFAULT_CACHET_PNG    = Path(__file__).parent / "templates" / "prescription" / "cachet.png"
DEFAULT_SIGNATURE_PNG = Path(__file__).parent / "templates" / "prescription" / "signature.png"
DEFAULT_LOGO_PNG      = Path(__file__).parent / "templates" / "prescription" / "logo.png"

AMY_TABLE: Dict[str, Dict[str, str]] = {
    'AMY 8':    {'label': "Mesure de l'acuité visuelle et de la réfraction – Renouvellement", 'price': '20,80 €'},
    'AMY 15':   {'label': 'Bilan des troubles oculomoteurs', 'price': '39,00 €'},
    'AMY 7,7':  {'label': 'Séance orthoptique (courte)', 'price': '15,40 €'},
    'AMY 7':    {'label': "Traitement de l'amblyopie par série de vingt séances", 'price': '18,20 €'},
    'AMY 4':    {'label': 'Traitement des hétérophories (20 séances)', 'price': '10,40 €'},
    'AMY 7.7':  {'label': 'Traitement du strabisme par série de vingt séances', 'price': '20,02 €'},
    'AMY 30':   {'label': 'Bilan orthoptique des déficiences visuelles (basse vision)', 'price': '78,00 €'},
    'AMY 30,5': {'label': 'Bilan des conséquences neuro-ophtalmologiques', 'price': '79,30 €'},
    'AMY 10':   {'label': 'Bilan des déséquilibres de la vision binoculaire', 'price': '26,00 €'},
    'AMY 14,5': {'label': 'Bilan des déséquilibres avec trouble neurosensoriel/accommodatif', 'price': '37,70 €'},
    'AMY 15,5': {'label': "Bilan d'une amblyopie", 'price': '40,30 €'},
    'AMY 19,2': {'label': 'Rééducation déficience visuelle - Plus de 16 ans (45 mn)', 'price': '49,92 €'},
    'AMY 13,2': {'label': 'Rééducation déficience visuelle - 16 ans et moins (30 mn)', 'price': '34,32 €'},
    'AMY 8,7':  {'label': 'Mesure acuité visuelle et réfraction - Primo-prescription', 'price': '22,62 €'},
}


def normalize_amy_code(raw: str) -> Optional[str]:
    if not raw:
        return None
    match = re.search(r'AMY\s*\(?\s*(\d+(?:[,\.]\d+)?)\s*\)?', raw, re.IGNORECASE)
    if not match:
        return None
    num = match.group(1).replace('.', ',')
    return f'AMY {num}'


def _format_date_fr(iso_date: Optional[str]) -> str:
    if not iso_date:
        return ""
    try:
        dt = datetime.strptime(iso_date, "%Y-%m-%d")
        return dt.strftime("%d/%m/%Y")
    except ValueError:
        return iso_date


def _find_para(paras, *needles, start: int = 0) -> int:
    """Return index of first paragraph (>= start) whose text contains all needles. -1 if none."""
    for i in range(start, len(paras)):
        text = paras[i].text or ""
        if all(n in text for n in needles):
            return i
    return -1


def _resolve_output_dir(template_path: Path) -> Path:
    """Pick a writable output dir for filled docs.

    Order:
      1. NEUROX_GENERATED_DIR env var
      2. %PROGRAMDATA%\\NeuroX\\generated  (production install)
      3. <template_dir>/filled              (dev fallback)
      4. system temp dir                    (last resort)
    """
    candidates = []
    env_dir = os.environ.get("NEUROX_GENERATED_DIR")
    if env_dir:
        candidates.append(Path(env_dir))
    program_data = os.environ.get("PROGRAMDATA")
    if program_data:
        candidates.append(Path(program_data) / "NeuroX" / "generated")
    candidates.append(template_path.parent / "filled")
    import tempfile
    candidates.append(Path(tempfile.gettempdir()) / "NeuroX" / "generated")

    for d in candidates:
        try:
            d.mkdir(parents=True, exist_ok=True)
            test = d / ".write_test"
            test.write_text("ok", encoding="utf-8")
            test.unlink()
            return d
        except OSError as e:
            logger.debug(f"Output dir candidate {d} unusable: {e}")
            continue
    # Should never reach here
    return Path(tempfile.gettempdir())


def _insert_acts(doc, acts_with_labels: List[Dict[str, str]]):
    actes_idx = None
    for i, p in enumerate(doc.paragraphs):
        if 'ACTES PRESCRITS' in p.text:
            actes_idx = i
            break
    if actes_idx is None:
        return
    for j, act in enumerate(acts_with_labels):
        target_idx = actes_idx + 1 + j
        if target_idx >= len(doc.paragraphs):
            break
        p = doc.paragraphs[target_idx]
        for run in p.runs:
            run.text = ""
        run = p.add_run(f"- {act['code']} : {act['label']}    ({act['price']})")
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def _replace_p28_with_table(doc, p28, cachet_path: Optional[str], signature_path: Optional[str]):
    """Replace P[28] with borderless 1x2 table: Cachet | Signature"""
    tbl = doc.add_table(rows=1, cols=2)

    for cell in tbl.rows[0].cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'nil')
            tcBorders.append(border)
        tcPr.append(tcBorders)

    for cell in tbl.rows[0].cells:
        cell.width = Cm(8)

    cell_left  = tbl.rows[0].cells[0]
    cell_right = tbl.rows[0].cells[1]

    p_left = cell_left.paragraphs[0]
    label_run = p_left.add_run("Cachet du centre\n")
    label_run.bold = True
    label_run.font.size = Pt(9)
    if cachet_path and Path(cachet_path).exists():
        try:
            p_left.add_run().add_picture(cachet_path, width=Inches(1.1))
        except Exception as e:
            logger.warning(f"Cachet error: {e}")
    else:
        p_left.add_run("[cachet manquant]")

    p_right = cell_right.paragraphs[0]
    label_run2 = p_right.add_run("Signature du prescripteur\n")
    label_run2.bold = True
    label_run2.font.size = Pt(9)
    if signature_path and Path(signature_path).exists():
        try:
            p_right.add_run().add_picture(signature_path, width=Inches(1.1))
        except Exception as e:
            logger.warning(f"Signature error: {e}")
    else:
        p_right.add_run("[signature manquante]")

    p28._element.addnext(tbl._tbl)

    for run in p28.runs:
        run.text = ""
    p = p28._p
    for r in p.findall(qn('w:r')):
        p.remove(r)


def fill_prescription_template(
    extraction_data: Dict,
    center_info: Optional[Dict[str, str]] = None,
    finess: str = "",
    city: str = "",
    amy_table: Optional[Dict[str, Dict[str, str]]] = None,
    template_path: Optional[str] = None,
    cachet_png: Optional[str] = None,
    signature_png: Optional[str] = None,
    logo_png: Optional[str] = None,
) -> str:

    tpl = Path(template_path) if template_path else DEFAULT_TEMPLATE
    if not tpl.exists():
        raise FileNotFoundError(f"Template not found: {tpl}")

    cachet_path    = cachet_png    or (str(DEFAULT_CACHET_PNG)    if DEFAULT_CACHET_PNG.exists()    else None)
    signature_path = signature_png or (str(DEFAULT_SIGNATURE_PNG) if DEFAULT_SIGNATURE_PNG.exists() else None)
    logo_path      = logo_png      or (str(DEFAULT_LOGO_PNG)      if DEFAULT_LOGO_PNG.exists()      else None)

    doc   = Document(str(tpl))
    paras = doc.paragraphs

    patient   = extraction_data.get("patient", {})
    doctor    = extraction_data.get("doctor", {})
    care      = extraction_data.get("orthoptic_care", {})
    form_date = _format_date_fr(extraction_data.get("form_date"))

    # --- Centre name + logo (template merges them on the same line) ---
    idx_centre = _find_para(paras, "Nom du centre")
    if idx_centre < 0:
        idx_centre = _find_para(paras, "centre")
    if idx_centre < 0 and paras:
        idx_centre = 0
    if idx_centre >= 0 and idx_centre < len(paras):
        p_centre = paras[idx_centre]
        # Replace centre name in the first non-empty run that contains "centre"
        if center_info and center_info.get("name"):
            replaced = False
            for run in p_centre.runs:
                if "centre" in run.text.lower() or "Nom du centre" in run.text:
                    run.text = center_info["name"]
                    replaced = True
                    break
            if not replaced and p_centre.runs:
                # fall back: replace first run with text
                for run in p_centre.runs:
                    if run.text.strip():
                        run.text = center_info["name"]
                        break
        # Replace "logos" placeholder with actual logo image
        if logo_path and Path(logo_path).exists():
            logo_run = None
            for run in p_centre.runs:
                if "logo" in run.text.lower():
                    logo_run = run
                    break
            try:
                if logo_run is not None:
                    logo_run.text = ""
                    logo_run.add_picture(logo_path, width=Inches(1.0))
                else:
                    # No "logos" placeholder: append picture as new run
                    p_centre.add_run().add_picture(logo_path, width=Inches(1.0))
                logger.info(f"Logo inserted: {logo_path}")
            except Exception as e:
                logger.warning(f"Logo error: {e}")
        else:
            # Strip "logos" placeholder text
            for run in p_centre.runs:
                if "logo" in run.text.lower():
                    run.text = ""

    # --- Adresse / Tél / Email ---
    idx_addr = _find_para(paras, "Adresse")
    if idx_addr >= 0 and center_info:
        p_addr = paras[idx_addr]
        for run in p_addr.runs:
            if "Adresse" in run.text:
                run.text = f"Adresse : {center_info.get('address', '')}"
            elif "Tél" in run.text or "T\u00e9l" in run.text:
                run.text = f"\nTél : {center_info.get('tel', '')}"
            elif "Email" in run.text:
                run.text = f"\nEmail : {center_info.get('email', '')}"
    elif idx_addr < 0:
        logger.debug("Adresse paragraph not found in template")

    # --- FINESS + Fait à ---
    idx_finess = _find_para(paras, "FINESS")
    if idx_finess >= 0:
        for run in paras[idx_finess].runs:
            if "FINESS" in run.text:
                run.text = f"FINESS : {finess}"
            elif "Fait" in run.text:
                run.text = f"Fait à {city}"
    else:
        logger.debug("FINESS paragraph not found in template")

    # --- Date DD / MM / YYYY ---
    if form_date:
        # Date paragraph: contains "Date" and "/" but NOT "naissance"
        idx_date = -1
        for i, p in enumerate(paras):
            t = p.text or ""
            if "Date" in t and "/" in t and "naissance" not in t.lower():
                idx_date = i
                break
        if idx_date >= 0:
            parts = form_date.split("/")
            if len(parts) == 3:
                runs = paras[idx_date].runs
                # Look for the three runs that come after a "/" separator pattern.
                # Strategy: find runs whose text is a separator '/' or whitespace, fill the empty/whitespace runs around them.
                # Fall back to the legacy fixed-index pattern if 6+ runs are present.
                if len(runs) >= 6:
                    # Heuristic: locate runs that contain only whitespace or are empty between '/' runs
                    slash_positions = [k for k, r in enumerate(runs) if r.text.strip() == "/"]
                    if len(slash_positions) >= 2:
                        # day before first '/', month between '/'s, year after last '/'
                        try:
                            day_run_idx   = slash_positions[0] - 1
                            month_run_idx = slash_positions[0] + 1
                            year_run_idx  = slash_positions[1] + 1
                            if 0 <= day_run_idx < len(runs):
                                runs[day_run_idx].text = parts[0]
                            if 0 <= month_run_idx < len(runs):
                                runs[month_run_idx].text = parts[1]
                            if 0 <= year_run_idx < len(runs):
                                runs[year_run_idx].text = parts[2]
                        except Exception as e:
                            logger.debug(f"Date run-fill heuristic failed: {e}")
                    else:
                        # legacy fixed-index fallback
                        runs[1].text = parts[0]
                        runs[3].text = parts[1]
                        runs[5].text = parts[2]
        else:
            logger.debug("Date paragraph not found in template")

    # --- Nom / Prénom ---
    last_name  = patient.get("last_name") or ""
    first_name = patient.get("first_name") or ""
    idx_nom = -1
    for i, p in enumerate(paras):
        t = p.text or ""
        if "Nom :" in t and ("Prénom" in t or "Pr\u00e9nom" in t):
            idx_nom = i
            break
    if idx_nom >= 0:
        for run in paras[idx_nom].runs:
            if "Nom :" in run.text:
                run.text = f"Nom : {last_name}"
            elif "Prénom :" in run.text or "Pr\u00e9nom :" in run.text:
                run.text = f"Prénom : {first_name}"
    else:
        logger.debug("Nom/Prénom paragraph not found in template")

    # --- NIR ---
    nir = patient.get("nir") or ""
    idx_nir = _find_para(paras, "Sociale")
    if idx_nir < 0:
        idx_nir = _find_para(paras, "NIR")
    if idx_nir >= 0:
        for run in paras[idx_nir].runs:
            if "Sociale" in run.text or "NIR" in run.text:
                run.text = f"N° Sécurité Sociale (NIR) : {nir}"
                break

    # --- Date de naissance ---
    birth = _format_date_fr(patient.get("birth_date"))
    idx_birth = _find_para(paras, "naissance")
    if idx_birth >= 0:
        for run in paras[idx_birth].runs:
            if "naissance" in run.text:
                run.text = f"Date de naissance : {birth}"
                break

    # --- Docteur / RPPS ---
    doc_name = doctor.get("full_name") or ""
    doc_name = re.sub(r'\s+[A-Z]$', '', doc_name).strip(' ,')
    rpps = doctor.get("rpps") or ""
    idx_doc = _find_para(paras, "Docteur")
    if idx_doc >= 0:
        runs = paras[idx_doc].runs
        if len(runs) >= 4:
            runs[0].text = f"Je soussigné(e), Docteur : {doc_name}"
            runs[1].text = "     "
            runs[2].text = "RPPS"
            runs[3].text = f" : {rpps}"
            # Clear any extra runs
            for extra in runs[4:]:
                extra.text = ""
        else:
            for run in runs:
                if "Docteur" in run.text:
                    run.text = f"Je soussigné(e), Docteur : {doc_name}"
                elif "RPPS" in run.text:
                    run.text = f"RPPS : {rpps}"

    # --- Les soins orthoptiques ---
    description = care.get("description") or ""
    idx_soins = _find_para(paras, "soins orthoptiques")
    if idx_soins >= 0:
        for run in paras[idx_soins].runs:
            if "soins orthoptiques" in run.text:
                run.text = f"Les soins orthoptiques suivants : {description}"
                break

    # --- Actes prescrits ---
    acts_raw = care.get("acts_prescribed", [])
    if acts_raw:
        lookup = amy_table if amy_table is not None else AMY_TABLE
        acts_with_labels = []
        for code in acts_raw:
            normalized = normalize_amy_code(code) or code
            info = lookup.get(normalized, {})
            acts_with_labels.append({
                "code": normalized,
                "label": info.get("label", ""),
                "price": info.get("price", ""),
            })
        _insert_acts(doc, acts_with_labels)

    # --- Cachet + Signature ---
    idx_cachet = -1
    for i, p in enumerate(paras):
        t = p.text or ""
        if ("Cachet" in t) or ("Signature" in t):
            idx_cachet = i
            break
    if idx_cachet >= 0:
        try:
            _replace_p28_with_table(doc, paras[idx_cachet], cachet_path, signature_path)
        except Exception as e:
            logger.warning(f"Cachet/Signature insertion failed: {e}")
    else:
        logger.warning("Cachet/Signature anchor paragraph not found — skipping")

    # --- Save ---
    output_dir = _resolve_output_dir(tpl)
    timestamp  = datetime.now().strftime("%Y%m%d_%H%M%S")
    filled_path = output_dir / f"Ordonnance_filled_{timestamp}.docx"
    doc.save(str(filled_path))
    logger.info(f"Filled template saved: {filled_path}")
    return str(filled_path)


def fill_and_convert_to_pdf(
    extraction_data: Dict,
    output_pdf_path: str,
    center_info: Optional[Dict[str, str]] = None,
    finess: str = "",
    city: str = "",
    amy_table: Optional[Dict[str, Dict[str, str]]] = None,
    template_path: Optional[str] = None,
    cachet_png: Optional[str] = None,
    signature_png: Optional[str] = None,
    logo_png: Optional[str] = None,
) -> str:

    filled_docx = fill_prescription_template(
        extraction_data=extraction_data,
        center_info=center_info,
        finess=finess,
        city=city,
        amy_table=amy_table,
        template_path=template_path,
        cachet_png=cachet_png,
        signature_png=signature_png,
        logo_png=logo_png,
    )

    os.makedirs(os.path.dirname(output_pdf_path), exist_ok=True)

    try:
        from docx2pdf import convert
        convert(filled_docx, output_pdf_path)
        logger.info(f"PDF generated: {output_pdf_path}")
    except Exception as e:
        logger.error(f"docx2pdf conversion failed: {e}")
        raise

    try:
        os.unlink(filled_docx)
    except OSError:
        pass

    return output_pdf_path